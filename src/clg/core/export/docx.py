"""DOCX renderer — builds a Word document with python-docx."""

from __future__ import annotations

import io

from clg.core.export.base import ExportError, ExportFormat, LetterDocument, split_name_header


class DocxRenderer:
    fmt = ExportFormat.docx
    media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    extension = "docx"

    def render(self, doc: LetterDocument) -> bytes:
        try:
            import docx
            from docx.shared import Pt

            document = docx.Document()
            # The content already includes its header and signature. The first
            # line is the candidate's name — render it bold and larger for
            # emphasis, then the rest of the letter as normal paragraphs.
            name, body = split_name_header(doc.content)
            if name:
                heading = document.add_paragraph()
                run = heading.add_run(name)
                run.bold = True
                run.font.size = Pt(20)
            # Split on blank lines into paragraphs; preserve single newlines
            # within a block (e.g. the contact header) as Word line breaks.
            for block in (b for b in body.split("\n\n") if b.strip()):
                paragraph = document.add_paragraph()
                for i, line in enumerate(block.split("\n")):
                    if i:
                        paragraph.add_run().add_break()
                    paragraph.add_run(line)
            buffer = io.BytesIO()
            document.save(buffer)
        except Exception as exc:  # noqa: BLE001 - normalize docx errors
            raise ExportError(f"DOCX rendering failed: {exc}") from exc
        return buffer.getvalue()
