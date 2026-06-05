"""Export: all five renderers produce well-formed, non-empty output."""

from __future__ import annotations

import io

from clg.core.export.base import ExportError, LetterDocument, render

DOC = LetterDocument(
    content=(
        "Vitor Cavalcante\n"
        "vitor@example.com · +55 61 99999-0000 · linkedin.com/in/vitor · Brasília, Brazil\n"
        "RE: Senior Engineer\n"
        "Dear Hiring Manager,\n\n"
        "I am excited to apply.\n\n"
        "Sincerely,\nVitor Cavalcante"
    ),
    title="Cover Letter",
    candidate_name="Vitor Cavalcante",
    language="en",
)


def test_text_formats():
    txt = render(DOC, "txt").decode()
    assert b"Dear Hiring Manager" in render(DOC, "txt")
    # TXT has no typography — the content (name header included) is verbatim.
    assert txt.count("vitor@example.com") == 1
    assert txt.startswith("Vitor Cavalcante\nvitor@example.com")

    html = render(DOC, "html").lower()
    assert b"<!doctype html>" in html


def test_name_header_is_emphasized():
    """The candidate's name (first line) is promoted to an emphasized heading."""
    # Markdown: name becomes an H1 (renders large + bold); body follows.
    md = render(DOC, "markdown").decode()
    assert md.startswith("# Vitor Cavalcante\n")
    assert "vitor@example.com" in md
    # The name must not be duplicated into the body block.
    assert "# Vitor Cavalcante\n\nvitor@example.com" in md

    # HTML: name rendered as <h1 class="name">; body paragraphs have no <h1>.
    html = render(DOC, "html").decode()
    assert '<h1 class="name">Vitor Cavalcante</h1>' in html
    assert "font-size: 22pt" in html
    assert html.count("<h1") == 1

    # DOCX: first paragraph is a bold, enlarged run with the name.
    import docx
    from docx.shared import Pt

    document = docx.Document(io.BytesIO(render(DOC, "docx")))
    first_run = document.paragraphs[0].runs[0]
    assert first_run.text == "Vitor Cavalcante"
    assert first_run.bold is True
    assert first_run.font.size == Pt(20)


def test_docx_is_a_zip():
    assert render(DOC, "docx")[:2] == b"PK"


def test_pdf_or_clear_error():
    """PDF renders when WeasyPrint libs exist; otherwise raises a clear error."""
    try:
        out = render(DOC, "pdf")
    except ExportError as exc:
        assert "WeasyPrint" in str(exc) or "PDF" in str(exc)
    else:
        assert out[:4] == b"%PDF"
