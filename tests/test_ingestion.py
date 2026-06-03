"""Ingestion: manual normalization, low-confidence signal, type/size guards."""

from __future__ import annotations

import io
from collections.abc import Callable

import pytest

from clg.core.ingestion.parse import (
    MAX_INPUT_BYTES,
    IngestionError,
    extract_upload,
    normalize_manual,
)


def _docx_bytes(build: Callable[[object], None]) -> bytes:
    import docx

    document = docx.Document()
    build(document)
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def test_manual_high_confidence():
    result = normalize_manual("a" * 100)
    assert result.low_confidence is False
    assert result.source == "manual"


def test_manual_low_confidence_on_short_text():
    assert normalize_manual("short").low_confidence is True


def test_unsupported_type_rejected():
    with pytest.raises(IngestionError):
        extract_upload("resume.txt", b"hello")


def test_oversized_pdf_rejected():
    with pytest.raises(IngestionError):
        extract_upload("resume.pdf", b"x" * (MAX_INPUT_BYTES + 1))


def test_docx_reads_paragraphs_and_tables():
    def build(doc):
        doc.add_paragraph("Jane Developer — Staff Engineer")
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Email"
        table.cell(0, 1).text = "jane@example.com"
        table.cell(1, 0).text = "Skills"
        table.cell(1, 1).text = "Python, Go, Kubernetes"

    result = extract_upload("cv.docx", _docx_bytes(build))

    # Body paragraph AND table cell content must survive (tables were dropped before).
    assert "Jane Developer" in result.text
    assert "jane@example.com" in result.text
    assert "Python, Go, Kubernetes" in result.text
    assert result.low_confidence is False


def test_docx_whitespace_is_normalized():
    def build(doc):
        doc.add_paragraph("Led    a    team   of   six   engineers")

    result = extract_upload("cv.docx", _docx_bytes(build))

    assert "Led a team of six engineers" in result.text
